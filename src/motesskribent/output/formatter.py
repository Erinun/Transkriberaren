"""Formatera transkribering till Markdown och JSON."""

from __future__ import annotations

import json
from datetime import datetime

from motesskribent.transcription.transcriber import TranscribedSegment, clean_transcription_text  # noqa: F401


def format_timestamp(seconds: float) -> str:
    """Formatera sekunder till HH:MM:SS (om >= 1 timme) eller MM:SS."""
    seconds = max(0.0, seconds)
    total_secs = int(seconds)
    hours = total_secs // 3600
    minutes = (total_secs % 3600) // 60
    secs = total_secs % 60

    if hours > 0:
        return f"{hours:02d}:{minutes:02d}:{secs:02d}"
    return f"{minutes:02d}:{secs:02d}"


def merge_short_segments(
    segments: list[TranscribedSegment],
    max_gap: float = 2.0,
) -> list[TranscribedSegment]:
    """Slå ihop konsekutiva segment från samma talare med gap < max_gap.

    Konkatenerar .text med mellanslag och mergar .words-listor.
    None speaker_id == None behandlas som samma talare.
    """
    if not segments:
        return []

    merged: list[TranscribedSegment] = [
        TranscribedSegment(
            text=segments[0].text,
            start=segments[0].start,
            end=segments[0].end,
            words=list(segments[0].words),
            speaker_id=segments[0].speaker_id,
            speaker_label=segments[0].speaker_label,
        )
    ]

    for seg in segments[1:]:
        prev = merged[-1]
        gap = seg.start - prev.end

        if seg.speaker_id == prev.speaker_id and gap < max_gap:
            prev.text = f"{prev.text} {seg.text}"
            prev.end = seg.end
            prev.words = prev.words + list(seg.words)
        else:
            merged.append(TranscribedSegment(
                text=seg.text,
                start=seg.start,
                end=seg.end,
                words=list(seg.words),
                speaker_id=seg.speaker_id,
                speaker_label=seg.speaker_label,
            ))

    return merged


def to_markdown(
    segments: list[TranscribedSegment],
    metadata: dict,
) -> str:
    """Generera svenskt mötesprotokoll i Markdown-format.

    metadata bör innehålla: date, duration, num_speakers,
    processing_time, model_name, version.
    """
    lines: list[str] = []

    lines.append("# Mötesprotokoll")
    lines.append("")

    date_str = metadata.get("date", datetime.now().strftime("%Y-%m-%d"))
    duration = metadata.get("duration", 0.0)
    num_speakers = metadata.get("num_speakers", 0)

    lines.append(f"**Datum:** {date_str}  ")
    lines.append(f"**Längd:** {format_timestamp(duration)}  ")
    lines.append(f"**Antal talare:** {num_speakers}  ")
    audio_source = metadata.get("audio_source")
    if audio_source:
        lines.append(f"**Ljudkälla:** {audio_source}  ")
    lines.append("")
    lines.append("---")
    lines.append("")

    for seg in segments:
        timestamp = format_timestamp(seg.start)
        label = seg.speaker_label or seg.speaker_id or "Okänd talare"
        lines.append(f"**[{timestamp}] {label}:**  ")
        lines.append(seg.text)
        lines.append("")

    lines.append("---")
    lines.append("")

    processing_time = metadata.get("processing_time", 0.0)
    model_name = metadata.get("model_name", "")
    version = metadata.get("version", "")

    lines.append(f"*Genererat av MötesSkribent {version}*  ")
    lines.append(f"*Modell: {model_name}*  ")
    lines.append(f"*Bearbetningstid: {format_timestamp(processing_time)}*  ")

    return "\n".join(lines) + "\n"


def to_json(
    segments: list[TranscribedSegment],
    metadata: dict,
    include_word_timestamps: bool = False,
) -> str:
    """Generera JSON med metadata, speakers-summary och segments.

    Returnerar json.dumps(ensure_ascii=False, indent=2).
    """
    unique_speakers: dict[str | None, str | None] = {}
    for seg in segments:
        if seg.speaker_id not in unique_speakers:
            unique_speakers[seg.speaker_id] = seg.speaker_label

    speakers_summary = [
        {"id": sid, "label": label}
        for sid, label in unique_speakers.items()
    ]

    segments_data = []
    for seg in segments:
        seg_dict: dict = {
            "start": round(seg.start, 3),
            "end": round(seg.end, 3),
            "speaker_id": seg.speaker_id,
            "speaker_label": seg.speaker_label,
            "text": seg.text,
        }

        if include_word_timestamps and seg.words:
            seg_dict["words"] = [
                {
                    "word": w.word,
                    "start": round(w.start, 3),
                    "end": round(w.end, 3),
                    "confidence": round(w.confidence, 4),
                }
                for w in seg.words
            ]

        segments_data.append(seg_dict)

    output = {
        "metadata": metadata,
        "speakers": speakers_summary,
        "segments": segments_data,
    }

    return json.dumps(output, ensure_ascii=False, indent=2)


def to_docx(
    segments: list[TranscribedSegment],
    metadata: dict,
    output_path: str,
) -> None:
    """Generera Word-dokument (.docx) med mötesprotokoll.

    Skriver direkt till angiven sökväg.
    """
    from docx import Document
    from docx.shared import Pt

    doc = Document()

    doc.add_heading("Mötesprotokoll", level=1)

    date_str = metadata.get("date", datetime.now().strftime("%Y-%m-%d"))
    duration = metadata.get("duration", 0.0)
    num_speakers = metadata.get("num_speakers", 0)

    meta_para = doc.add_paragraph()
    meta_para.add_run(f"Datum: {date_str}\n").bold = True
    meta_para.add_run(f"Längd: {format_timestamp(duration)}\n")
    meta_para.add_run(f"Antal talare: {num_speakers}\n")
    audio_source = metadata.get("audio_source")
    if audio_source:
        meta_para.add_run(f"Ljudkälla: {audio_source}\n")

    doc.add_paragraph("_" * 50)

    for seg in segments:
        timestamp = format_timestamp(seg.start)
        label = seg.speaker_label or seg.speaker_id or "Okänd talare"

        p = doc.add_paragraph()
        run = p.add_run(f"[{timestamp}] {label}:")
        run.bold = True
        run.font.size = Pt(11)

        text_para = doc.add_paragraph(seg.text)
        text_para.runs[0].font.size = Pt(11)

    doc.add_paragraph("_" * 50)

    processing_time = metadata.get("processing_time", 0.0)
    model_name = metadata.get("model_name", "")
    version = metadata.get("version", "")

    footer = doc.add_paragraph()
    footer.add_run(f"Genererat av MötesSkribent {version}\n").italic = True
    footer.add_run(f"Modell: {model_name}\n").italic = True
    footer.add_run(f"Bearbetningstid: {format_timestamp(processing_time)}").italic = True

    doc.save(output_path)
